import docx
import nltk
import pymorphy2
import tkinter as tk
import tkinter.filedialog as fd

def read_doc_file(file_path):
    if file_path.endswith('.docx'):
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    elif file_path.endswith('.doc'):
        with open(file_path, 'rb') as file:
            doc = docx.Document(file)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return '\n'.join(full_text)
    else:
        print("what the fuck")
        return "Unsupported file format"


def get_lexemes(text):
    # разбиваем текст на слова
    words = nltk.word_tokenize(text.lower(), language='russian')
    # создаем объект MorphAnalyzer для морфологического анализа слов
    morph = pymorphy2.MorphAnalyzer()
    lexemes = []
    for word in words:
        # пропускаем знаки препинания
        if not word.isalpha():
            continue
        # получаем морфологические характеристики слова
        parse = morph.parse(word)[0]
        # создаем запись для каждой лексемы
        lexeme = {
            'base': parse.normal_form,  # основа слова
            'pos': parse.tag.POS,  # часть речи
            'gender': parse.tag.gender,  # род
            'number': parse.tag.number,  # число
            'case': parse.tag.case,  # падеж
            'tense': parse.tag.tense,  # время (для глаголов)
            'aspect': parse.tag.aspect  # вид (для глаголов)
        }
        # добавляем лексему в список
        lexemes.append(lexeme)
    # сортируем список лексем по алфавиту
    lexemes.sort(key=lambda x: x['base'])
    return lexemes


def generate_word_form(word: str, gender: str, case: str, number: str):
    """
    Функция генерирует словоформу по заданным параметрам
    word: базовая форма слова
    gender: род: "муж", "жен", "сред"
    case: падеж: "им", "род", "дат", "вин", "твор", "пр", "мест"
    number: число: "ед", "мн"
    return: словоформа в соответствии с заданными параметрами
    """
    morph = pymorphy2.MorphAnalyzer()
    parse = morph.parse(word)[0]
    tags = {"муж": "masc", "жен": "femn", "сред": "neut", "им": "nomn", "род": "gent",
            "дат": "datv", "вин": "accs", "твор": "ablt", "пр": "loct","ед": "sing", "мн": "plur"}
    gram_info = {tags[v] for v in {gender, case, number}}
    print('THIS ',gram_info)
    return parse.inflect(gram_info).word

def get_lexems_to_text(lexems):
    txt = ""
    tags = {"NOUN":"сущ.", "ADJF":"прилаг.", "ADJS":"прилаг.", "COMP":"компаратив", "VERB":"глагол", "INFN":"глагол","PRTF":"причастие",
            "PRTS":"причастие", "GRND":"дееприч.", "NUMR":"числит.", "ADVB": "наречие", "NRPO":"местоим.-сущ.","PRED":"предикатив",
            "PREP":"предлог","CONJ":"союз","PRCL":"частица","INTJ":"междометие","nomn":"Им.п.", "gent":"Род.п.","datv":"Дат.п.",
            "accs":"Вин.п.","ablt":"Твор.п.","loct":"Пред.п.","sing":"ед.ч.","plur":"мн.ч","masc":"муж.р.","femn":"жен.р.","neut":"ср.р"}
    for item in lexems:
        txt += "Основа: "
        txt += item['base'] + ", ч.речи: "
        txt += tags[item['pos']]
        if item['gender'] != None:
            txt +=", род: " + tags[item['gender']]
        if item['number'] !=None:
            txt +=', число: ' + tags[item['number']]
        if item['case'] != None:
            txt +=', падеж: ' + tags[item['case']]

        txt +='\n'
    return txt




# txt = read_doc_file('TEST.docx')
# print(txt)
# lst = get_lexemes("Объем работы и прочие требования")
# print(lst)
# print(get_lexems_to_text(lst))
#
# print(generate_word_form("стул", "муж", "род", "ед"))  # "стула"
# print(generate_word_form("стул", "муж", "дат", "ед"))  # "стулу"
# print(generate_word_form("стул", "муж", "твор", "мн"))  # "стульях"
#
def generate_bttn():
    word.delete(1.0,tk.END)
    base = input1_entry.get()
    gender = input2_entry.get()
    case = input3_entry.get()
    numb = input4_entry.get()
    word.insert(tk.END, generate_word_form(base,gender,case,numb))
    input1_entry.delete(0,tk.END)
    input2_entry.delete(0,tk.END)
    input3_entry.delete(0,tk.END)
    input4_entry.delete(0,tk.END)




def open_file():
    filetypes = (
        ('DOCX', '*.docx'),
        ('DOC', '*.doc')
    )
    filename = fd.askopenfilename(
        title='Open a file',
        initialdir='/',
        filetypes=filetypes
    )
    if filename:
        doc = docx.Document(filename)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        text = get_lexems_to_text(get_lexemes(text))
        new_window = tk.Toplevel(root)
        new_window.title(filename)
        new_window.geometry('400x300')
        text_widget = tk.Text(new_window)
        text_widget.insert('1.0', text)
        text_widget.pack()

root = tk.Tk()
root.title('DOC Reader')
root.geometry('400x300')
open_button = tk.Button(root, text='Открыть файл doc или docx', command=open_file)
open_button.pack()

input1_label = tk.Label(root, text='Основа слова:')
input1_label.pack()
input1_entry = tk.Entry(root)
input1_entry.pack()

input2_label = tk.Label(root, text='Род:')
input2_label.pack()
input2_entry = tk.Entry(root)
input2_entry.pack()

input3_label = tk.Label(root, text='Падеж:')
input3_label.pack()
input3_entry = tk.Entry(root)
input3_entry.pack()

input4_label = tk.Label(root, text='Число:')
input4_label.pack()
input4_entry = tk.Entry(root)
input4_entry.pack()

generate_button = tk.Button(root, text='Генерация слова', command=generate_bttn)
generate_button.pack()

word = tk.Text(root)
word.pack()

root.mainloop()